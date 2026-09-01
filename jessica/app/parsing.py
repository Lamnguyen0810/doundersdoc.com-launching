"""
Turn an uploaded contract into an ordered list of clauses:
[{"ref": "2.1", "heading": "Equity grant", "text": "..."}]
DOCX keeps real paragraph structure; PDF is text-extracted and split heuristically.
"""
import io
import re

import pymupdf
from docx import Document as DocxDocument

CLAUSE_RE = re.compile(r"^\s*((?:\d+\.)+\d*|\d+|[A-Z]\.|\([a-z]\)|\([ivx]+\))\s+(.*)$")
HEADING_MAX = 90


def parse(data: bytes, mime: str, filename: str) -> tuple[list[dict], int | None]:
    name = filename.lower()
    if name.endswith(".docx") or "wordprocessingml" in mime:
        return _parse_docx(data), None
    if name.endswith(".pdf") or mime == "application/pdf":
        return _parse_pdf(data)
    if name.endswith(".txt") or mime.startswith("text/"):
        return _split_paragraphs(data.decode("utf-8", errors="replace").splitlines()), None
    raise ValueError("Unsupported file type. Upload a .docx, .pdf or .txt file.")


def _parse_docx(data: bytes) -> list[dict]:
    doc = DocxDocument(io.BytesIO(data))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return _split_paragraphs(lines)


def _parse_pdf(data: bytes) -> tuple[list[dict], int]:
    lines: list[str] = []
    with pymupdf.open(stream=data, filetype="pdf") as pdf:
        pages = pdf.page_count
        for page in pdf:
            text = page.get_text("text")
            lines.extend(ln.strip() for ln in text.splitlines() if ln.strip())
    return _split_paragraphs(lines), pages


def _split_paragraphs(lines: list[str]) -> list[dict]:
    clauses: list[dict] = []
    current: dict | None = None
    seq = 0
    for line in lines:
        m = CLAUSE_RE.match(line)
        if m:
            if current:
                clauses.append(current)
            ref, rest = m.group(1).rstrip("."), m.group(2).strip()
            heading = rest if len(rest) <= HEADING_MAX and not rest.endswith(".") else ""
            current = {"ref": ref, "heading": heading, "text": "" if heading else rest}
        elif current is not None:
            current["text"] = (current["text"] + " " + line).strip()
        else:
            seq += 1
            current = {"ref": f"P{seq}", "heading": "", "text": line}
    if current:
        clauses.append(current)
    return clauses


def clauses_to_text(clauses: list[dict]) -> str:
    out = []
    for c in clauses:
        head = f" {c['heading']}" if c.get("heading") else ""
        out.append(f"[{c['ref']}]{head}\n{c['text']}".strip())
    return "\n\n".join(out)
