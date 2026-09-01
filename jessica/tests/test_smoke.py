import io

from docx import Document

from app.llm.rubric import RUBRIC, compute_grade
from app.parsing import parse


def make_docx() -> bytes:
    d = Document()
    d.add_paragraph("ADVISER AGREEMENT")
    d.add_paragraph("1. Appointment")
    d.add_paragraph("The Company appoints the Adviser to provide strategic advice.")
    d.add_paragraph("2. Equity")
    d.add_paragraph("2.1 The Adviser shall be granted options over 1% of the fully diluted share capital.")
    d.add_paragraph("2.2 The options vest monthly over 24 months with no cliff.")
    d.add_paragraph("7. Termination")
    d.add_paragraph("Either party may terminate on 30 days' written notice.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_parse_docx_into_clauses():
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    clauses, pages = parse(make_docx(), mime, "a.docx")
    refs = [c["ref"] for c in clauses]
    assert "2.1" in refs and "2.2" in refs and "7" in refs
    assert pages is None
    two_one = next(c for c in clauses if c["ref"] == "2.1")
    assert "1%" in two_one["text"]


def test_grade_is_deterministic_and_bounded():
    perfect = {r["key"]: 4 for r in RUBRIC}
    zero = {r["key"]: 0 for r in RUBRIC}
    assert compute_grade(perfect) == (100, "A")
    assert compute_grade(zero) == (0, "E")
    mid = {r["key"]: 2 for r in RUBRIC}
    assert compute_grade(mid)[0] == 50


def test_health_and_home(client):
    assert client.get("/health").json() == {"ok": True}
    r = client.get("/")
    assert r.status_code == 200
    assert "Review a contract" in r.text


def test_upload_creates_document_and_job(client):
    files = {"file": ("adviser.docx", make_docx(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = client.post("/documents", files=files, data={"perspective": "the company"}, follow_redirects=False)
    assert r.status_code == 303
    location = r.headers["location"]
    assert location.startswith("/documents/") and "job=" in location
    page = client.get(location)
    assert page.status_code == 200 and "adviser.docx" in page.text
    job_id = location.split("job=")[1]
    status = client.get(f"/jobs/{job_id}")
    assert status.status_code == 200 and "Queued" in status.text
